import io
import base64
from pathlib import Path

import docx
from PIL import Image

from core.config import settings
from core.logger import logger
from rag.parser.utils import ParsedChunk, table_to_markdown, split_text


def parse_docx(file_path: str | Path) -> list[ParsedChunk]:
    source = Path(file_path).name
    chunks: list[ParsedChunk] = []
    doc = docx.Document(str(file_path))

    # Text — paragraphs
    text_blocks = [p.text for p in doc.paragraphs if p.text.strip()]
    combined = "\n".join(text_blocks)
    if combined.strip():
        for i, chunk_text in enumerate(split_text(combined, settings.chunk_size, settings.chunk_overlap)):
            chunks.append(ParsedChunk(
                text=chunk_text,
                type="text",
                page=0,
                source=source,
                metadata={"chunk_index": i},
            ))

    # Tables — preserve headers, stored as atomic markdown chunks
    for table_idx, table in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        markdown, headers = table_to_markdown(rows)
        if markdown.strip():
            chunks.append(ParsedChunk(
                text=markdown,
                type="table",
                page=0,
                source=source,
                metadata={"table_index": table_idx, "headers": headers},
            ))

    # Images — embedded blobs via relationships
    img_idx = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                pil_image = Image.open(io.BytesIO(rel.target_part.blob)).convert("RGB")
                buf = io.BytesIO()
                pil_image.save(buf, format="PNG")
                image_b64 = base64.b64encode(buf.getvalue()).decode()
                chunks.append(ParsedChunk(
                    text=f"[Image {img_idx} in document]",
                    type="image",
                    page=0,
                    source=source,
                    metadata={"image_index": img_idx, "image_base64": image_b64},
                ))
                img_idx += 1
            except Exception as e:
                logger.warning(f"Skipped image {img_idx}: {e}")

    logger.info(f"Parsed DOCX {source}: {len(chunks)} chunks")
    return chunks
