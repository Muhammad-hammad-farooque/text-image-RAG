import os
import base64
import io
import numpy as np
import torch
import fitz
import docx

from dotenv import load_dotenv
from PIL import Image
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chat_models import init_chat_model
from transformers import CLIPModel, CLIPProcessor

load_dotenv()

# Load CLIP once globally
clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")
clip_model.eval()

splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
llm = init_chat_model("meta-llama/llama-4-scout-17b-16e-instruct", model_provider="groq")


# ── Embedding ─────────────────────────────────────────────────────────────────

def embed_image(image_data):
    image = Image.open(image_data).convert("RGB") if isinstance(image_data, str) else image_data
    inputs = clip_processor(images=image, return_tensors="pt")
    with torch.no_grad():
        vision_outputs = clip_model.vision_model(pixel_values=inputs["pixel_values"])
        features = clip_model.visual_projection(vision_outputs.pooler_output)
        features = features / features.norm(dim=-1, keepdim=True)
        return features.squeeze().numpy()


def embed_text(text):
    inputs = clip_processor(text=text, return_tensors="pt", padding=True, truncation=True, max_length=77)
    with torch.no_grad():
        text_outputs = clip_model.text_model(
            input_ids=inputs["input_ids"],
            attention_mask=inputs["attention_mask"]
        )
        features = clip_model.text_projection(text_outputs.pooler_output)
        features = features / features.norm(dim=-1, keepdim=True)
        return features.squeeze().numpy()


# ── Document processors ───────────────────────────────────────────────────────

def process_pdf(file_path, all_docs, all_embeddings, image_data_stores):
    doc = fitz.open(file_path)
    for i, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            chunks = splitter.split_documents([Document(page_content=text, metadata={"page": i, "type": "text"})])
            for chunk in chunks:
                all_embeddings.append(embed_text(chunk.page_content))
                all_docs.append(chunk)

        for img_index, img in enumerate(page.get_images(full=True)):
            try:
                base_image = doc.extract_image(img[0])
                pil_image = Image.open(io.BytesIO(base_image["image"])).convert("RGB")
                image_id = f"page_{i}_img_{img_index}"
                buffered = io.BytesIO()
                pil_image.save(buffered, format="PNG")
                image_data_stores[image_id] = base64.b64encode(buffered.getvalue()).decode()
                all_embeddings.append(embed_image(pil_image))
                all_docs.append(Document(
                    page_content=f"[Image:{image_id}]",
                    metadata={"page": i, "type": "image", "image_id": image_id}
                ))
            except Exception:
                continue
    doc.close()


def process_docx(file_path, all_docs, all_embeddings, image_data_stores):
    doc = docx.Document(file_path)

    text_blocks = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                text_blocks.append(row_text)

    combined_text = "\n".join(text_blocks)
    if combined_text.strip():
        chunks = splitter.split_documents([Document(page_content=combined_text, metadata={"page": 0, "type": "text"})])
        for i, chunk in enumerate(chunks):
            chunk.metadata["page"] = i
            all_embeddings.append(embed_text(chunk.page_content))
            all_docs.append(chunk)

    img_index = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                pil_image = Image.open(io.BytesIO(rel.target_part.blob)).convert("RGB")
                image_id = f"docx_img_{img_index}"
                buffered = io.BytesIO()
                pil_image.save(buffered, format="PNG")
                image_data_stores[image_id] = base64.b64encode(buffered.getvalue()).decode()
                all_embeddings.append(embed_image(pil_image))
                all_docs.append(Document(
                    page_content=f"[Image:{image_id}]",
                    metadata={"page": img_index, "type": "image", "image_id": image_id}
                ))
                img_index += 1
            except Exception:
                continue


def process_txt(file_path, all_docs, all_embeddings):
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    if text.strip():
        chunks = splitter.split_documents([Document(page_content=text, metadata={"page": 0, "type": "text"})])
        for i, chunk in enumerate(chunks):
            chunk.metadata["page"] = i
            all_embeddings.append(embed_text(chunk.page_content))
            all_docs.append(chunk)


# ── Public API ────────────────────────────────────────────────────────────────

def load_document(file_path):
    """Process a document and return a FAISS vector store + image store."""
    all_docs, all_embeddings, image_data_stores = [], [], {}

    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        process_pdf(file_path, all_docs, all_embeddings, image_data_stores)
    elif ext == ".docx":
        process_docx(file_path, all_docs, all_embeddings, image_data_stores)
    elif ext == ".txt":
        process_txt(file_path, all_docs, all_embeddings)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt")

    embeddings_array = np.array(all_embeddings)
    vector_store = FAISS.from_embeddings(
        text_embeddings=[(doc.page_content, emb) for doc, emb in zip(all_docs, embeddings_array)],
        embedding=None,
        metadatas=[doc.metadata for doc in all_docs]
    )

    return vector_store, image_data_stores, len(all_docs), len(image_data_stores)


def retrieve(query, vector_store, k=5):
    """Retrieve top-k relevant documents for a query."""
    query_embedding = embed_text(query)
    return vector_store.similarity_search_by_vector(embedding=query_embedding, k=k)


def build_message(query, retrieved_docs, image_data_stores):
    """Build a multimodal HumanMessage from retrieved docs."""
    content = [{"type": "text", "text": f"Question: {query}\n\nContext:\n"}]

    text_docs = [d for d in retrieved_docs if d.metadata.get("type") == "text"]
    image_docs = [d for d in retrieved_docs if d.metadata.get("type") == "image"]

    if text_docs:
        text_context = "\n\n".join(
            f"[Page {d.metadata['page']}]: {d.page_content}" for d in text_docs
        )
        content.append({"type": "text", "text": f"Text excerpts:\n{text_context}\n"})

    for doc in image_docs:
        image_id = doc.metadata.get("image_id")
        if image_id and image_id in image_data_stores:
            content.append({"type": "text", "text": f"\n[Image from page {doc.metadata['page']}]:\n"})
            content.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{image_data_stores[image_id]}"}
            })

    content.append({"type": "text", "text": "\n\nPlease answer the question based on the provided text and images."})
    return HumanMessage(content=content)
